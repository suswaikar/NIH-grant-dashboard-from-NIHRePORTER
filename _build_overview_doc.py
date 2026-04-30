"""Build Evans Endowment Awards Overview document for the Research Advisory Committee."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path(r"C:\Users\swaikar\OneDrive - Boston University\Department of Medicine"
           r"\Chair advisory committee for Evans\Evans_Endowment_Awards_Overview.docx")

doc = Document()

# --- Styles ---
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level, size, color in [(1, 16, "1F4E79"), (2, 13, "2E75B6"), (3, 11, "2E75B6")]:
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Arial"
    hs.font.size = Pt(size)
    hs.font.color.rgb = RGBColor.from_string(color)
    hs.font.bold = True

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = "Arial"
    # Data rows
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Arial"
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

# ============================================================================
# TITLE
# ============================================================================
title = doc.add_heading("Evans Endowment: Research Awards Overview", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph("Department of Medicine Research Advisory Committee")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(11)
subtitle.runs[0].font.color.rgb = RGBColor.from_string("666666")
subtitle.runs[0].italic = True

doc.add_paragraph("April 2026", style="Normal").alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============================================================================
# OVERVIEW
# ============================================================================
doc.add_heading("Overview", level=2)

doc.add_paragraph(
    "The Evans Endowment supports early-career and mid-career faculty through four primary "
    "award mechanisms: K award salary gap support, Pilot awards, Junior Faculty awards, and "
    "Bridge funding. This report summarizes the Department of Medicine's investment across "
    "these programs and evaluates their return on investment through subsequent NIH funding."
)

# ============================================================================
# KEY METRICS
# ============================================================================
doc.add_heading("Investment Summary", level=2)

doc.add_paragraph(
    "From FY2016 to FY2026, the Department has invested a total of $6.4 million across "
    "all four award programs, supporting 89 unique investigators."
)

add_table(doc,
    ["Program", "Period", "Investigators", "Total Investment"],
    [
        ["K Award Salary Gap", "FY2016-FY2026", "59", "$4,379,357"],
        ["Junior Faculty Award", "FY2022-FY2026", "12", "$1,200,000"],
        ["Pilot Award", "FY2022-FY2026", "18", "$497,078"],
        ["Bridge Funding", "FY2025-FY2026", "10", "$304,000"],
        ["Total", "FY2016-FY2026", "89*", "$6,380,435"],
    ],
    col_widths=[2.2, 1.5, 1.2, 1.5],
)

p = doc.add_paragraph()
run = p.add_run("*Some investigators received multiple award types. GT97 awards are excluded "
                "as they are an accounting mechanism, not a research investment.")
run.font.size = Pt(8)
run.italic = True
run.font.color.rgb = RGBColor.from_string("888888")

# ============================================================================
# K AWARD RESULTS
# ============================================================================
doc.add_heading("K Award Program: K-to-R Transition", level=2)

doc.add_paragraph(
    "The K award salary gap program has supported 59 investigators across 11 DoM sections "
    "over 11 fiscal years. The program covers 50% of the unfunded salary gap plus fringe "
    "benefits (28.8%) for junior faculty with career development awards."
)

doc.add_heading("Key Findings", level=3)

findings = [
    "83% of K awardees remain at CAMED/BMC (49 of 59), demonstrating strong retention.",
    "51% are female (30F, 29M), reflecting balanced sex representation in the pipeline.",
    "Among eligible K awardees (those whose K support ended before FY2025, allowing sufficient "
    "time to secure subsequent funding), approximately 59% have secured post-K NIH funding "
    "(R01 or equivalent). Awardees still receiving K support or whose support ended within "
    "the last year are excluded from this calculation.",
    "K awardees who transitioned have collectively generated over $60 million in post-K "
    "NIH direct costs — a return of approximately 14x on the Department's $4.4 million investment.",
    "Two K awardees have advanced to the rank of Professor, 14 to Associate Professor, "
    "and 3 have moved to industry leadership positions.",
]
for f in findings:
    doc.add_paragraph(f, style="List Bullet")

doc.add_heading("K Awards by Section", level=3)

add_table(doc,
    ["Section", "Awardees", "DoM Investment", "Post-K NIH Direct"],
    [
        ["Cardiovascular Medicine", "6", "$488,139", "High"],
        ["Pulmonary", "15", "$580,000+", "High"],
        ["Infectious Diseases", "11", "$474,000+", "High"],
        ["General Internal Medicine", "11", "$568,000+", "Moderate"],
        ["Nephrology", "5", "$216,000+", "Moderate"],
        ["Other Sections", "11", "$300,000+", "Varies"],
    ],
    col_widths=[2.2, 1.0, 1.5, 1.7],
)

p = doc.add_paragraph()
run = p.add_run("Note: Post-K NIH Direct costs are conservative estimates based on NIH RePORTER "
                "data. Only grants newly initiated after the K support period are counted. "
                "Exact figures available in the interactive dashboard.")
run.font.size = Pt(8)
run.italic = True
run.font.color.rgb = RGBColor.from_string("888888")

# ============================================================================
# PILOT / JR FACULTY
# ============================================================================
doc.add_heading("Pilot & Junior Faculty Awards", level=2)

doc.add_paragraph(
    "Since FY2022, the Department has awarded $497,078 in pilot grants to 18 investigators "
    "and $1,200,000 in junior faculty awards to 12 investigators. These programs provide "
    "seed funding to generate preliminary data for competitive NIH applications."
)

doc.add_paragraph(
    "Early evidence suggests strong returns: several pilot and junior faculty awardees have "
    "subsequently secured R01-equivalent NIH funding. Detailed ROI analysis for these newer "
    "programs is available in the interactive dashboard, though the shorter time horizon "
    "(FY2022 onward) limits the number of awardees who have had sufficient time to convert "
    "their pilot data into funded grants."
)

# ============================================================================
# BRIDGE FUNDING
# ============================================================================
doc.add_heading("Bridge Funding", level=2)

doc.add_paragraph(
    "Bridge funding was introduced in FY2025 to support investigators between grant cycles. "
    "To date, 10 awards totaling $304,000 have been made (3 in FY2025, 7 in FY2026). "
    "This program addresses a critical gap period that can derail productive research programs "
    "and is too new to evaluate for ROI."
)

# ============================================================================
# SEX BREAKDOWN
# ============================================================================
doc.add_heading("Biological Sex", level=2)

doc.add_paragraph(
    "Across all K awardees, representation is balanced: 30 female (51%) and 29 male (49%). "
    "K-to-R transition rates and post-K NIH funding are comparable across sex, though sample "
    "sizes for eligible transitioners are small. Detailed breakdowns by sex are available "
    "in the interactive dashboard."
)

# ============================================================================
# DASHBOARD
# ============================================================================
doc.add_heading("Interactive Dashboard", level=2)

doc.add_paragraph(
    "A live interactive dashboard has been developed to enable real-time exploration of "
    "all award data, individual investigator NIH grant portfolios, and ROI metrics. "
    "The dashboard queries NIH RePORTER in real time and is accessible via a password-protected "
    "Streamlit Cloud deployment."
)

features = [
    "Awards Overview: Investment table and trend chart across all programs and fiscal years",
    "ROI Summary: K-to-R transition analysis and Pilot/Junior Faculty ROI, with expandable "
    "section-level detail showing individual awardees",
    "Sex Breakdown: Awards and conversion rates by biological sex",
    "Drill Down: View individual awardees for any award type and fiscal year",
    "Investigator Lookup: Select any awardee to see their complete NIH grant portfolio "
    "from RePORTER, separated into K grants, post-K grants, and all grants",
]
for f in features:
    doc.add_paragraph(f, style="List Bullet")

# ============================================================================
# METHODOLOGY
# ============================================================================
doc.add_heading("Methodology Notes", level=2)

methods = [
    "NIH grant data is sourced from the NIH RePORTER API v2 (public, queried in real time).",
    "K-to-R transition rate uses a conservative definition: only non-K NIH grants with fiscal "
    "year strictly after the last year of K support are counted as transitions.",
    "A 1-year buffer is applied: awardees whose K ended in FY2025 or later are excluded from "
    "the transition rate denominator, as they have not had sufficient time to secure subsequent funding.",
    "K24 awards (midcareer mentoring) are treated as successful outcomes, not filtered out.",
    "GT97 awards are excluded from all ROI calculations.",
    "Name matching against NIH RePORTER uses first 3 characters of first name + full last name. "
    "Known false positives (common names) are manually filtered.",
]
for m in methods:
    doc.add_paragraph(m, style="List Bullet")

# ============================================================================
# SAVE
# ============================================================================
doc.save(str(OUT))
print(f"Saved to: {OUT}")
